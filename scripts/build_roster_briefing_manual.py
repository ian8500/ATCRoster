from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("deliverables/ATC_Roster_and_Briefing_User_Manual.docx")
BLUE = "1F4E79"
DARK_BLUE = "17365D"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4CE"
PALE_GREEN = "E2F0D9"
PALE_RED = "FCE4D6"
INK = "172B4D"
MUTED = "5B6573"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, PALE_BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = str(value)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, label, text, tone="blue"):
    fills = {"blue": PALE_BLUE, "gold": PALE_GOLD, "green": PALE_GREEN, "red": PALE_RED}
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fills[tone])
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "1")
        edge.set(qn("w:space"), "5")
        edge.set(qn("w:color"), fills[tone])
        borders.append(edge)
    p_pr.append(borders)
    lead = p.add_run(f"{label}: ")
    lead.bold = True
    lead.font.color.rgb = RGBColor.from_string(DARK_BLUE if tone != "red" else "9B1C1C")
    p.add_run(text)


def add_bullets(doc, items, level=0):
    for text in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(text)


def new_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend((tabs, indent, spacing))
    level.extend((start, num_fmt, level_text, suffix, p_pr))
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_steps(doc, steps):
    num_id = new_decimal_numbering(doc)
    for title, detail in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_element = OxmlElement("w:numId")
        num_id_element.set(qn("w:val"), str(num_id))
        num_pr.extend((ilvl, num_id_element))
        p_pr.append(num_pr)
        r = p.add_run(title)
        r.bold = True
        p.add_run(f" — {detail}")


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    return p


def add_definition(doc, label, detail):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(detail)


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, DARK_BLUE),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.375)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.188)
    styles["List Bullet 2"].paragraph_format.left_indent = Inches(0.65)
    styles["List Bullet 2"].paragraph_format.first_line_indent = Inches(-0.188)
    styles["List Number"].paragraph_format.left_indent = Inches(0.375)
    styles["List Number"].paragraph_format.first_line_indent = Inches(-0.188)


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = "ATC Roster  |  Roster & Briefing User Manual"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(footer)

    # Editorial cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(110)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("USER MANUAL")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("ATC Roster and Briefing")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(70)
    r = p.add_run("A practical guide for operational staff, roster editors and unit administrators")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    add_callout(
        doc,
        "Purpose",
        "Use this manual to complete day-to-day roster and briefing work safely, consistently and with a clear audit trail.",
        "blue",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(45)
    p.add_run("Version 1.0  |  30 July 2026").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Prepared for customer demonstration and controlled operational rollout")

    add_page_break(doc)
    add_heading(doc, "Document control", 1)
    add_table(
        doc,
        ["Field", "Detail"],
        [
            ("Document", "ATC Roster and Briefing User Manual"),
            ("Version", "1.0"),
            ("Date", "30 July 2026"),
            ("Audience", "Operational staff, roster editors, unit administrators and management"),
            ("System scope", "Roster and Briefing modules"),
            ("Review trigger", "Material workflow change, new release or local operating-procedure change"),
        ],
        [2700, 6660],
    )
    add_callout(
        doc,
        "Local procedures",
        "This manual explains how to use the software. Your unit’s approved operational, fatigue, briefing and record-retention procedures always take precedence.",
        "gold",
    )

    add_heading(doc, "How to use this manual", 2)
    add_bullets(doc, [
        "Operational staff should begin with Sections 1–4 and Sections 9–11.",
        "Roster editors should also read Sections 5–8.",
        "Briefing publishers and administrators should also read Sections 12–16.",
        "Unit administrators should use the launch and routine-control checklists in Sections 17–18.",
    ])

    add_heading(doc, "Contents", 1)
    contents = [
        "1. System overview and access",
        "2. Roles, permissions and responsibilities",
        "3. Navigation and common controls",
        "4. Reading the monthly roster",
        "5. Editing and annotating the roster",
        "6. Fatigue warnings and roster publication",
        "7. Leave, sickness, requests and overtime",
        "8. Roster reports, exports and routine controls",
        "9. Briefing module overview",
        "10. Staff briefing workflow",
        "11. Reading, acknowledging and archiving",
        "12. Briefing settings and message types",
        "13. Publishing instructions and Briefs of the Day",
        "14. Briefing register and withdrawal",
        "15. Briefing reports and assurance",
        "16. Briefing audit evidence",
        "17. Troubleshooting and support",
        "18. Launch and daily operating checklists",
        "Appendix A. Quick-reference workflows",
        "Appendix B. Glossary",
    ]
    for item in contents:
        doc.add_paragraph(item)

    add_page_break(doc)
    add_heading(doc, "1. System overview and access", 1)
    doc.add_paragraph(
        "ATC Roster provides a shared operational workspace. The Roster module records duties, watches, availability and associated staffing information. The Briefing module distributes controlled instructions and temporary daily messages, records access and acknowledgement, and provides management oversight."
    )
    add_heading(doc, "Signing in", 2)
    add_steps(doc, [
        ("Open the customer system address", "Use the address supplied by your unit administrator."),
        ("Enter your username and password", "Usernames are issued by the unit; passwords are case-sensitive."),
        ("Select Sign in", "After successful authentication, the module selection page opens."),
        ("Select a module", "Choose Roster or Briefing. Select Home at any time to return to the module selection page."),
    ])
    add_callout(
        doc,
        "Security",
        "Never share an account. Acknowledgements, edits and audit events are attributed to the signed-in user. Sign out or lock the device when leaving it unattended.",
        "red",
    )
    add_heading(doc, "Password and username recovery", 2)
    doc.add_paragraph(
        "The login page provides separate links for a forgotten password and a forgotten username. Follow the displayed recovery process. If recovery is unavailable or your account is locked, contact your unit administrator."
    )
    add_heading(doc, "Signing out", 2)
    doc.add_paragraph(
        "Open the main menu and select Logout. On a shared device, close the browser after the logout confirmation appears."
    )

    add_heading(doc, "2. Roles, permissions and responsibilities", 1)
    doc.add_paragraph(
        "The options visible to you depend on your account role and the permissions configured by your unit. Missing controls normally indicate that your account does not have permission for that action."
    )
    add_table(
        doc,
        ["Role", "Typical access", "Primary responsibility"],
        [
            ("Operational user", "Published/read-only roster, own requests, assigned briefings and personal archive", "Check duties; read and acknowledge instructions"),
            ("Roster editor", "Roster editing and operational roster tools permitted by the unit", "Maintain accurate assignments and annotations"),
            ("Administrator", "Publishing, configuration, reports, audit and account administration", "Control publication and oversee compliance"),
            ("Management / authorised reviewer", "Reports and assurance according to assigned permissions", "Review trends, exceptions and evidence"),
        ],
        [1800, 3780, 3780],
    )
    add_callout(
        doc,
        "Access denied",
        "A 403 page means that you are authenticated but not authorised for that area. Use the module return link and ask a Unit Administrator if your duties require additional access.",
        "gold",
    )

    add_heading(doc, "3. Navigation and common controls", 1)
    add_definition(doc, "Home", "Returns to the module selection page.")
    add_definition(doc, "Unit name", "Remains at the top left so the active operational unit is always clear.")
    add_definition(doc, "Menu", "On smaller screens, opens the navigation links. On a desktop, navigation may already be visible.")
    add_definition(doc, "Status messages", "Success, warning and error messages appear near the top of the page after an action.")
    add_definition(doc, "Confirmation prompts", "Destructive or material actions, such as withdrawal or unpublishing, require confirmation.")
    add_heading(doc, "Using a phone or tablet", 2)
    add_bullets(doc, [
        "Use Menu to show or hide navigation.",
        "The monthly roster is wider than a phone screen by design. Swipe horizontally inside the roster frame.",
        "Use Fit width for an overview and 75%, 90% or 100% for progressively larger text and cells.",
        "For detailed roster editing or long document review, a desktop or larger tablet is recommended.",
    ])

    add_page_break(doc)
    add_heading(doc, "Part I — Roster module", 1)
    add_heading(doc, "4. Reading the monthly roster", 1)
    doc.add_paragraph(
        "The roster opens on a monthly grid. Staff appear in rows and calendar dates appear in columns. The left-hand columns identify the person, staff number, medical and unit-endorsement dates, and watch."
    )
    add_heading(doc, "Move between months", 2)
    add_steps(doc, [
        ("Select the previous-month button", "Moves one calendar month backwards."),
        ("Select the next-month button", "Moves one calendar month forwards."),
        ("Confirm the title", "Always check the month and year before reading or changing a duty."),
    ])
    add_heading(doc, "Understand a roster cell", 2)
    add_table(
        doc,
        ["Element", "Meaning"],
        [
            ("Duty code", "The roster code assigned for that person and date."),
            ("Blank / dash", "No duty code is currently assigned."),
            ("Annotation", "Additional operational information attached to the duty."),
            ("Warning symbol", "One or more configured fatigue checks require attention."),
            ("Request marker", "A request or proposed change is associated with the date."),
            ("Today highlight", "Identifies the current calendar day."),
        ],
        [2500, 6860],
    )
    add_heading(doc, "Expiry and training colours", 2)
    add_table(
        doc,
        ["Colour", "Meaning"],
        [
            ("Green", "More than 90 days remain."),
            ("Amber", "90 days or fewer remain."),
            ("Red", "The recorded date has expired."),
            ("Orange / UT", "The person is recorded as under training."),
        ],
        [2500, 6860],
    )
    add_callout(
        doc,
        "Important",
        "A coloured date is a visual prompt, not a substitute for local licence or competency control. Report incorrect dates immediately.",
        "gold",
    )
    add_heading(doc, "Draft and published status", 2)
    doc.add_paragraph(
        "The status panel below the grid identifies the month as Draft or Published. A draft may change. Publication marks the roster as the current issued version and notifies staff according to system configuration."
    )

    add_heading(doc, "5. Editing and annotating the roster", 1)
    doc.add_paragraph(
        "Only authorised users can edit roster cells. Other users see the same controls in a disabled, read-only state."
    )
    add_heading(doc, "Change a duty code", 2)
    add_steps(doc, [
        ("Confirm the staff member and date", "Read both the row label and column heading before making a change."),
        ("Open the duty-code list", "Only roster codes configured in the system are available."),
        ("Select the required code", "The system submits and validates the change."),
        ("Review the result", "Check the saved code, any warning message and any change to totals."),
    ])
    add_callout(
        doc,
        "Code control",
        "Free-text duty codes are not accepted. If the required code is missing, an administrator must configure the corresponding roster code before it can be used.",
        "blue",
    )
    add_heading(doc, "Add or update an annotation", 2)
    add_steps(doc, [
        ("Select the annotation control in the required cell", "Available annotations are limited to configured roster annotations."),
        ("Choose the appropriate annotation", "Use only the annotation that accurately describes the operational situation."),
        ("Add supporting text where offered", "Select the annotation badge to open its text editor, enter the detail and select Save text."),
        ("Remove an annotation when no longer applicable", "Use the remove control and confirm the cell is clear."),
    ])
    add_heading(doc, "Editing good practice", 2)
    add_bullets(doc, [
        "Make one logical change at a time and confirm it has saved.",
        "Do not use annotations as a substitute for an approved roster code.",
        "Recheck the affected staff member’s adjacent duties after a change.",
        "If a published roster must change, follow the unit’s change-control and notification procedure.",
    ])

    add_heading(doc, "6. Fatigue warnings and roster publication", 1)
    add_heading(doc, "Fatigue warnings", 2)
    doc.add_paragraph(
        "The system evaluates configured fatigue rules against working-duty segments and displays warnings on relevant working roster cells. Hover over or focus the warning to read the rule and calculated result."
    )
    add_steps(doc, [
        ("Read the complete warning", "Identify the rule code and the affected rest, duty or recovery value."),
        ("Review surrounding duties", "Check preceding and following assignments, including overnight duties."),
        ("Apply local escalation", "Do not dismiss a warning solely because the software permits the roster to be saved."),
        ("Record an authorised correction", "Amend the roster only when the correct operational decision is known."),
    ])
    add_callout(
        doc,
        "Safety-critical control",
        "Fatigue warnings support decision-making; they do not replace competent operational review, approved fatigue procedures or management responsibility.",
        "red",
    )
    add_heading(doc, "Publish a roster", 2)
    add_steps(doc, [
        ("Complete the month", "Resolve known gaps, invalid assignments and relevant warnings."),
        ("Check staff and watch coverage", "Confirm the correct people, watch allocations and special staffing requirements."),
        ("Check associated absence and requests", "Ensure approved leave, sickness and accepted requests are reflected."),
        ("Select Publish roster", "Read the confirmation message, then confirm publication."),
        ("Verify Published roster", "Confirm the published status and publication timestamp are displayed."),
    ])
    add_heading(doc, "Undo publication", 2)
    doc.add_paragraph(
        "Authorised users can select Undo publication. The roster returns to Draft and staff may be notified. Use this only under the unit’s roster change-control procedure; the action is not a silent edit."
    )

    add_heading(doc, "7. Leave, sickness, requests and overtime", 1)
    add_heading(doc, "Leave and sickness", 2)
    doc.add_paragraph(
        "Open Leave / Sickness to record and review absence information. Administrators may also maintain the permitted leave and sickness types."
    )
    add_steps(doc, [
        ("Choose the staff member", "Confirm the correct person before entering dates."),
        ("Choose the type", "Select the configured leave or sickness category."),
        ("Enter the start and end dates", "Check that the range is inclusive and correctly ordered."),
        ("Save the new record", "New entries use Save Leave or the corresponding sickness action."),
        ("Maintain an existing record", "Existing data presents Update and Delete actions."),
    ])
    add_callout(
        doc,
        "Sensitive information",
        "Record only the information needed for roster administration. Do not enter unnecessary medical details in free-text fields.",
        "gold",
    )
    add_heading(doc, "Requests", 2)
    doc.add_paragraph(
        "The Requests area lets staff submit the request types made available by their unit and lets authorised users review them. A request is not an approved roster change until it has completed the unit’s approval workflow and appears correctly on the roster."
    )
    add_heading(doc, "Overtime Finder", 2)
    doc.add_paragraph(
        "The Overtime Finder supports identification of potential availability against roster information. Treat results as decision support: verify qualifications, fatigue implications, availability and local approval before confirming overtime."
    )

    add_heading(doc, "8. Roster reports, exports and routine controls", 1)
    add_heading(doc, "Export CSV", 2)
    add_steps(doc, [
        ("Open the required roster month", "The export uses the month currently displayed."),
        ("Select Export CSV", "Your browser downloads a comma-separated file."),
        ("Store and share it appropriately", "The export may contain staff and operational information."),
    ])
    add_heading(doc, "Print", 2)
    doc.add_paragraph(
        "Select Print to open the browser’s print workflow. Check the month, orientation, scale and page preview before printing or saving as PDF."
    )
    add_heading(doc, "Reports and metrics", 2)
    doc.add_paragraph(
        "Authorised users can use Reports and Metrics for operational summaries, leave and sickness reporting, and exportable totals. Always confirm the selected date range and unit context before relying on a report."
    )
    add_heading(doc, "Recommended monthly roster control", 2)
    add_bullets(doc, [
        "Confirm active staff and watches.",
        "Check licence, medical and endorsement expiry prompts.",
        "Review duty-code validity and unfilled cells.",
        "Review fatigue warnings and special staffing requirements.",
        "Reconcile approved leave, sickness and requests.",
        "Publish only after an authorised final review.",
        "Export or print only where operationally required and store securely.",
    ])

    add_page_break(doc)
    add_heading(doc, "Part II — Briefing module", 1)
    add_heading(doc, "9. Briefing module overview", 1)
    doc.add_paragraph(
        "The Briefing module presents each user with current material selected for them and keeps administration separate from the roster page. The main areas are My briefing and Archive, with Publish, Reports, Audit and Settings available to authorised users."
    )
    add_table(
        doc,
        ["Area", "Purpose"],
        [
            ("My briefing", "Current Briefs of the Day, mandatory messages and other messages assigned to the signed-in user."),
            ("Archive", "The user’s acknowledged and archived instructions, grouped by message type."),
            ("Publish", "Create, publish, withdraw and review briefing material."),
            ("Reports", "Compare login and roster activity; identify unread mandatory and other instructions."),
            ("Audit", "Review append-only evidence of publication, access, acknowledgement and assurance activity."),
            ("Settings", "Maintain the instruction message types available to publishers."),
        ],
        [1900, 7460],
    )
    add_heading(doc, "Briefing types", 2)
    add_definition(doc, "Uploaded instruction", "A PDF or Word document with an effective period, audience and message type. It may require acknowledgement.")
    add_definition(doc, "Brief of the Day", "A temporary text message shown directly on the staff Briefing home page for its effective period.")
    add_definition(doc, "Mandatory message", "An instruction for which the recipient must confirm that it has been read and understood.")
    add_definition(doc, "Other message", "A current non-mandatory instruction assigned to the user.")

    add_heading(doc, "10. Staff briefing workflow", 1)
    add_heading(doc, "Start-of-duty check", 2)
    add_steps(doc, [
        ("Sign in with your own account", "The system records briefing activity against your identity."),
        ("Select Briefing", "The My briefing page opens."),
        ("Review Brief of the Day", "Expand longer text where required and read all active entries."),
        ("Review Mandatory Messages", "Open every outstanding mandatory instruction."),
        ("Review Other Messages", "Open new or changed information relevant to your duties."),
        ("Complete acknowledgements", "Confirm only after reading and understanding the material."),
        ("Escalate uncertainty", "Contact the nominated supervisor or publisher if the instruction is unclear or appears incorrect."),
    ])
    add_callout(
        doc,
        "Operational discipline",
        "Do not acknowledge on behalf of another person or before reading the complete instruction. An acknowledgement is recorded evidence linked to your account.",
        "red",
    )
    add_heading(doc, "Home-page categories", 2)
    doc.add_paragraph(
        "Each category shows its current count. If nothing is assigned, the section displays an up-to-date message. Items shown depend on audience targeting, publication status, effective time and expiry time."
    )

    add_heading(doc, "11. Reading, acknowledging and archiving", 1)
    add_heading(doc, "Open and read an instruction", 2)
    add_steps(doc, [
        ("Select the instruction card", "The reader page shows its type, mandatory status, version and effective period."),
        ("Review the document", "PDFs open in the embedded viewer with full-screen, pop-out and download controls. Word files open in the device’s document viewer."),
        ("Keep the page active while reading", "Active reading time pauses when the page is not active."),
        ("Check the version and expiry", "Ensure you are reading the current material."),
    ])
    add_heading(doc, "Acknowledge an instruction", 2)
    add_steps(doc, [
        ("Finish reading the content", "Do not use the acknowledgement as a simple dismissal control."),
        ("Select the confirmation checkbox", "This confirms that you have read and understood the briefing."),
        ("Select Acknowledge", "The system records the date and time."),
        ("Verify the acknowledgement", "A green acknowledged status and timestamp replace the form."),
    ])
    add_heading(doc, "Archive or delete from your personal view", 2)
    doc.add_paragraph(
        "After acknowledgement, use Archive to move an instruction to your personal archive. The Archive page groups items by message type and provides Review and Delete. Delete removes the item from your personal view; the system’s audit evidence is retained."
    )
    add_callout(
        doc,
        "Version awareness",
        "Reports record the version read. If a materially revised instruction is issued, treat the new version as a new operational reading requirement.",
        "blue",
    )

    add_heading(doc, "12. Briefing settings and message types", 1)
    doc.add_paragraph(
        "Before publishing uploaded instructions, an administrator must maintain the message types used in the instruction dropdown and in users’ archives."
    )
    add_steps(doc, [
        ("Open Briefing and select Settings", "Only authorised users see this option."),
        ("Review the current message types", "Use names that staff will recognise consistently."),
        ("Enter one type per line", "Avoid duplicates, abbreviations that could be misunderstood and temporary project labels."),
        ("Select Save message types", "Return to Publish and confirm the correct options are available."),
    ])
    add_bullets(doc, [
        "Examples may include Operational Instruction, Safety Notice, Technical Instruction or Management Notice.",
        "Use the unit’s controlled-document taxonomy where one exists.",
        "Do not remove or rename categories without considering existing archived material and reports.",
    ])

    add_heading(doc, "13. Publishing instructions and Briefs of the Day", 1)
    add_heading(doc, "Publish an uploaded instruction", 2)
    add_steps(doc, [
        ("Open Publish", "Confirm that Document storage reports a healthy status."),
        ("Select Uploaded instruction", "The document, message-type and mandatory fields become available."),
        ("Choose an instruction message type", "Use the approved category."),
        ("Enter a clear title", "Make it specific enough for staff to identify the subject and version."),
        ("Choose the PDF or Word file", "Only supported PDF and DOCX files should be uploaded."),
        ("Set effective and expiry date/time", "The expiry must be later than the effective time."),
        ("Choose the user group", "Select Everyone, All operational staff, selected watches, selected roles or named individuals."),
        ("Set Mandatory acknowledgement if required", "Apply this according to the instruction’s operational and governance requirements."),
        ("Save draft or Publish now", "Use draft for review; publish only when authorised."),
        ("Wait for completion", "Keep the page open while the secure upload or publishing progress panel is displayed."),
    ])
    add_heading(doc, "Publish a Brief of the Day", 2)
    add_steps(doc, [
        ("Select Brief of the day", "The file field is replaced by a text field."),
        ("Enter the title and briefing text", "Keep it concise, current and operationally clear."),
        ("Set the effective and expiry date/time", "Use a tightly controlled period appropriate to the message."),
        ("Choose the audience", "Check watches, roles or individuals carefully when using targeted scopes."),
        ("Save draft or Publish now", "Review spelling, dates and recipients before publication."),
    ])
    add_callout(
        doc,
        "Recipient check",
        "Targeting is part of the safety control. Confirm the intended watch, role or named recipients before publishing; do not assume Everyone is always appropriate.",
        "gold",
    )
    add_heading(doc, "Draft review checklist", 2)
    add_bullets(doc, [
        "Correct title, document and message type",
        "Correct version and approved content",
        "Effective and expiry times use the intended local operational time",
        "Correct audience and mandatory status",
        "No unnecessary personal or sensitive information",
        "Publisher has authority to issue the material",
    ])

    add_heading(doc, "14. Briefing register and withdrawal", 1)
    add_heading(doc, "Current instructions", 2)
    doc.add_paragraph(
        "The Briefing register shows current instructions with title, version, creator, type, effective period, status and available action. Drafts can be published; published items can be withdrawn."
    )
    add_heading(doc, "Older instructions", 2)
    doc.add_paragraph(
        "Expired or otherwise historic instructions appear under Older instructions. Expand the expiry year, then the expiry month, to locate the record. Categorisation is based on the instruction’s expiry date."
    )
    add_heading(doc, "Withdraw a published briefing", 2)
    add_steps(doc, [
        ("Locate the current item", "Confirm the title, version and effective period."),
        ("Select Withdraw", "Read the confirmation prompt."),
        ("Confirm withdrawal", "The instruction is no longer treated as currently published."),
        ("Record any operational follow-up", "Use the unit’s procedure to replace, correct or notify affected staff."),
    ])
    add_callout(
        doc,
        "Audit retention",
        "Withdrawal does not erase the history. Publication and withdrawal activity remains available in the audit evidence.",
        "blue",
    )

    add_heading(doc, "15. Briefing reports and assurance", 1)
    doc.add_paragraph(
        "Reports provide a real-time oversight view for a selected operational date and preserve previous report runs."
    )
    add_heading(doc, "Run a report", 2)
    add_steps(doc, [
        ("Open Reports", "Confirm you are in the Briefing module."),
        ("Choose the operational date", "Use the date relevant to the management or assurance check."),
        ("Select Run report", "The system creates and saves a new report run."),
        ("Expand each result group", "Review exceptions before relying on the summary."),
    ])
    add_table(
        doc,
        ["Report group", "What it shows", "How to use it"],
        [
            ("Login and roster activity", "Each user’s last login and last rostered working day, highlighting differences", "Identify accounts or roster activity requiring review"),
            ("On duty with unread mandatory messages", "On-duty users with mandatory items not opened or opened but not acknowledged", "Prioritise immediate operational follow-up"),
            ("Unread instructions by user", "Outstanding instructions, type, reading state and effective period", "Manage individual briefing completion"),
            ("Read instructions and active reading time", "Acknowledged versions, timestamps and active viewing duration", "Review evidence and unusual patterns"),
            ("Previous reports", "Saved report runs and headline findings", "Compare checks and retain management oversight"),
        ],
        [2100, 3630, 3630],
    )
    add_callout(
        doc,
        "Interpretation",
        "Reading time is supporting evidence, not proof of comprehension. Very short or long durations may justify review, but must be interpreted with document length, device behaviour and operational context.",
        "gold",
    )
    add_heading(doc, "Delete a saved report", 2)
    doc.add_paragraph(
        "Use Delete beside the relevant previous report and confirm the action. Deletion is irreversible, so apply the unit’s retention and authorisation rules."
    )

    add_heading(doc, "16. Briefing audit evidence", 1)
    doc.add_paragraph(
        "The Audit page presents append-only evidence of publication, access, acknowledgement and assurance activity. Each row includes the time, actor, event type, briefing identifier and event detail."
    )
    add_heading(doc, "Using the audit view", 2)
    add_steps(doc, [
        ("Identify the event under review", "Start with the approximate date, user and instruction."),
        ("Check the actor and event type", "Confirm who performed the action and what was recorded."),
        ("Check the briefing identifier and detail", "Use these fields to distinguish similar titles or versions."),
        ("Preserve evidence appropriately", "Follow the unit’s incident, audit and data-protection procedures."),
    ])
    add_bullets(doc, [
        "Do not treat the audit page as an editing interface.",
        "Do not disclose audit details to unauthorised people.",
        "Escalate unexpected events, missing evidence or apparent account misuse.",
    ])

    add_page_break(doc)
    add_heading(doc, "17. Troubleshooting and support", 1)
    add_table(
        doc,
        ["Symptom", "Likely cause", "Action"],
        [
            ("A module is not shown on Home", "The module is not enabled for the unit or your account lacks access", "Contact the Unit Administrator"),
            ("403 access page", "Your account role does not permit the action", "Return to the module and request the required permission"),
            ("Roster controls are disabled", "You have read-only access or the action is restricted", "Use Requests or contact an authorised roster editor"),
            ("Required roster code is missing", "No corresponding active roster code exists", "Ask an administrator to configure the code; do not substitute free text"),
            ("Unexpected fatigue warning", "The configured rule has identified a duty/rest condition", "Read the full warning and apply local fatigue review"),
            ("Instruction upload unavailable", "Briefing document storage is unhealthy", "Do not repeatedly upload; contact support or an administrator"),
            ("Briefing item is not visible", "It may be a draft, outside its effective period, expired or targeted to another audience", "Check status, dates and recipients"),
            ("PDF does not display", "Browser viewer or device restriction", "Use Pop out or Download; confirm the file opens securely"),
            ("Acknowledgement will not submit", "Confirmation was not selected or the session expired", "Select the confirmation; refresh and sign in again if needed"),
            ("500 / internal server error", "Unexpected application or service failure", "Record the time, page and reference; avoid repeating changes; contact support"),
        ],
        [2360, 3220, 3780],
    )
    add_heading(doc, "Information to provide when reporting a problem", 2)
    add_bullets(doc, [
        "Your unit and username (never your password)",
        "Date and local time of the problem",
        "Module, page and action attempted",
        "Displayed error text and reference identifier",
        "Roster month, staff member or briefing title involved",
        "Whether the problem occurs on another supported browser or device",
        "A screenshot only if it does not expose information to an unauthorised recipient",
    ])

    add_heading(doc, "18. Launch and daily operating checklists", 1)
    add_heading(doc, "Unit launch checklist", 2)
    add_bullets(doc, [
        "Confirm the unit name, timezone and operational context.",
        "Create and verify watches, staff accounts and staff numbers.",
        "Configure valid roster codes and shift times.",
        "Load medical, licence and endorsement dates and verify sample colour states.",
        "Assign roles and test one ordinary user, one roster editor and one administrator.",
        "Create a representative roster, test editing, fatigue warnings, publication, export and printing.",
        "Configure Briefing message types.",
        "Test a draft, published instruction, Brief of the Day, mandatory acknowledgement and withdrawal.",
        "Run a Briefing report and review audit evidence.",
        "Confirm document storage health, support contacts, backups and the approved launch decision.",
    ])
    add_heading(doc, "Roster editor pre-publication checklist", 2)
    add_bullets(doc, [
        "Correct month, staff and watch population",
        "No unintended blank or invalid duty cells",
        "Leave, sickness and approved requests reconciled",
        "Fatigue warnings reviewed under local procedures",
        "Special staffing requirements checked",
        "Licence, medical and endorsement prompts reviewed",
        "Authorised reviewer approval obtained",
        "Published status and timestamp verified",
    ])
    add_heading(doc, "Briefing publisher checklist", 2)
    add_bullets(doc, [
        "Approved source content and version",
        "Clear title and correct message type",
        "Correct effective and expiry time",
        "Correct audience and mandatory status",
        "Storage reports healthy",
        "Publication visible to a test recipient where appropriate",
        "Acknowledgement and report behaviour verified",
        "Superseded material withdrawn under change control",
    ])
    add_heading(doc, "Operational staff start-of-duty checklist", 2)
    add_bullets(doc, [
        "Sign in using your own account.",
        "Check the current published roster and any authorised changes.",
        "Review all active Briefs of the Day.",
        "Open and acknowledge mandatory instructions after reading and understanding them.",
        "Review other assigned messages.",
        "Escalate any uncertainty, missing information or incorrect assignment.",
    ])

    add_page_break(doc)
    add_heading(doc, "Appendix A — Quick-reference workflows", 1)
    add_heading(doc, "Publish a monthly roster", 2)
    doc.add_paragraph("Review month → reconcile duties and absence → review warnings → obtain approval → Publish roster → verify timestamp.")
    add_heading(doc, "Issue a mandatory instruction", 2)
    doc.add_paragraph("Settings/type check → Upload instruction → dates → audience → Mandatory → draft review → publish → staff acknowledge → run report.")
    add_heading(doc, "Issue a Brief of the Day", 2)
    doc.add_paragraph("Publish → Brief of the day → title and text → dates → audience → review → publish → verify staff display.")
    add_heading(doc, "Correct published material", 2)
    doc.add_paragraph("Confirm affected item → withdraw/unpublish under authority → correct source information → reissue → notify affected users → review audit/report evidence.")
    add_heading(doc, "Respond to unread mandatory material", 2)
    doc.add_paragraph("Run Briefing report → expand on-duty exceptions → identify not opened vs opened/not acknowledged → follow local escalation → rerun or record outcome.")

    add_heading(doc, "Appendix B — Glossary", 1)
    glossary = [
        ("Acknowledgement", "A user’s recorded confirmation that they have read and understood a briefing instruction."),
        ("Active reading time", "Time accumulated while the instruction reader page is active; it pauses when the page is inactive."),
        ("Brief of the Day", "A temporary text briefing displayed directly on the user’s Briefing home page."),
        ("Briefing register", "Administrative list of current and older briefing instructions."),
        ("Draft roster", "A roster month that has not been issued as the current published roster."),
        ("Effective period", "The date and time range during which a briefing is current."),
        ("Fatigue warning", "A system-generated finding against a configured fatigue rule."),
        ("Mandatory message", "An instruction that requires an acknowledgement from each recipient."),
        ("Published roster", "The roster month marked and issued as the current published version."),
        ("Roster code", "A configured code representing a duty or other roster state."),
        ("Target scope", "The audience selected for a briefing: everyone, operational staff, watches, roles or individuals."),
        ("UT", "Under training."),
        ("Withdrawal", "Removal of a published briefing from current circulation while retaining its audit history."),
    ]
    for term, meaning in glossary:
        add_definition(doc, term, meaning)

    add_heading(doc, "End of manual", 1)
    doc.add_paragraph(
        "For local access, workflow or data questions, contact your Unit Administrator. For application faults, use the approved customer support route and include the diagnostic information listed in Section 17."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "ATC Roster and Briefing User Manual"
    doc.core_properties.subject = "Operational user guide for the Roster and Briefing modules"
    doc.core_properties.author = "IDAviation"
    doc.core_properties.keywords = "ATC, roster, briefing, user manual"
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()
